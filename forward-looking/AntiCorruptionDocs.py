import json
from tkinter import *
from tksheet import Sheet
from tkinter.filedialog import *
from docx import Document
from docx.shared import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


def load_data(file_path: str) -> list:
    with open(file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
        return data if isinstance(data, list) else []


def dump_data(file_path: str, data: list) -> None:
    with open(file_path, 'w') as f:
        json.dump(data, f)
        return None


def startpage() -> int and str:

    root = Tk()

    global status
    status = 0

    global file_path_in
    file_path_in = ''

    with open('templates/firstpage.txt', 'r', encoding='utf-8') as f:
        info_text = f.read()

    def close_window():
        global status
        status = 1
        root.destroy()

    def load_new():
        global file_path_in
        file_path_in = 'templates/table_new.json'
        root.destroy()

    def load_old():
        global file_path_in
        file_path_in = 'users/table_old.json'
        root.destroy()

    def load_clinic_spb():
        global file_path_in
        file_path_in = 'templates/types/clinic_spb.json'
        root.destroy()

    root.overrideredirect(True)
    root.state('zoomed')
    root.option_add('*tearOff', FALSE)
    menu = Menu(root)
    root.config(menu=menu)
    filemenu = Menu(menu, font=('Arial', 14))
    menu.add_cascade(label='   = Выбрать вариант работы с картой рисков =   ',
                     menu=filemenu)
    filemenu.add_command(label='Создать пустую карту рисков',
                         command=load_new)
    filemenu.add_command(label='Загрузить сохраненную карту рисков',
                         command=load_old)
    filemenu.add_command(label='Загрузить рекомендуемую карту рисков '
                               '"Государственное учреждение здравоохранения '
                               'Санкт-Петербурга"',
                         command=load_clinic_spb)
    menu.add_command(label='   = Выйти =   ', command=close_window)

    Label(root, compound='left', justify='left', wraplength=0,
          anchor='e', font=('Arial', 20), text=info_text).pack()

    root.mainloop()

    return status, file_path_in


def firstpage(data: list) -> int:

    root = Tk()

    global data_current
    data_current = data

    global status
    status = 0

    def save_map():
        file_path_out = 'users/table_old.json'
        global data_current
        dump_data(file_path_out, data_current)
        root.destroy()

    def close_window():
        global status
        status = 1
        root.destroy()

    root.title('Заполнение карты коррупционных рисков '
               'и мер по их минимизации (устранению)')
    root.overrideredirect(True)
    root.state('zoomed')
    root.grid_columnconfigure(index=0, weight=1)
    root.grid_rowconfigure(index=0, weight=1)

    menu = Menu(root)
    root.config(menu=menu)
    menu.add_command(label='   = Сохранить карту рисков и перейти дальше =   ',
                     command=save_map)
    menu.add_command(label='   = Выйти без сохранения карты =   ',
                     command=close_window)

    frame = Frame(root)
    frame.grid_columnconfigure(index=0, weight=1)
    frame.grid_rowconfigure(index=0, weight=1)
    frame.grid(row=0, column=0, sticky='nswe')

    sheet = Sheet(parent=frame,
                  headers=data[0],
                  data=data[1:],
                  set_all_heights_and_widths=True,
                  max_column_width='500',
                  header_height='5',
                  header_align='left',
                  row_index_align='left')
    sheet.enable_bindings()
    sheet.grid(row=0, column=0, sticky='nswe')

    root.mainloop()

    return status


def secondpage() -> None:

    root = Tk()

    def save_client_params():
        for i in [0, 1]:
            for j in range(0, len(client_params[i])):
                elem_text = client_params[i][j]
                while '\n' in elem_text:
                    elem_text = elem_text.replace('\n', '')
                    client_params[i][j] = elem_text
        dump_data('users/client_params.json', client_params)
        return None

    def download_docs():

        def add_data(data: list) -> list:
            list_5 = ['низкий', 'средний', 'высокий']
            list_6 = ['низкая', 'средняя', 'высокая']
            data[0].insert(7, 'Уровень коррупционного риска')
            for i in range(1, len(data)):
                element_5 = data[i][5].lower().strip()
                element_6 = data[i][6].lower().strip()
                position_5 = [j for j, x in enumerate(list_5) if x == element_5]
                position_6 = [j for j, x in enumerate(list_6) if x == element_6]
                if position_5 and position_6:
                    sum_elements = int(position_5[0]) + int(position_6[0])
                    if sum_elements < 2:
                        element_7 = 'Низкий'
                    elif sum_elements == 2:
                        element_7 = 'Средний'
                    else:
                        element_7 = 'Высокий'
                else:
                    element_7 = ''
                data[i].insert(7, element_7)
            return data

        def text_replacement(text: str, data: list, client_params: list) -> list:

            index_begin = text.find('<index_')

            while index_begin != -1:
                index_begin = text.find('<index_')
                index_fin = text.find('>')
                doc_type = text[index_begin+7]

                if doc_type == '1':
                    index1 = int(text[index_begin+9])
                    index2 = text[(index_begin+11):index_fin]

                    if index2 == 'L':
                        positions = [i for i, x in enumerate(client_params[index1]) if x == 1]
                        if positions:
                            last_position = positions[-1]
                        else:
                            last_position = None
                        list_return = [int(doc_type),
                                       client_params[index1],
                                       last_position]
                        return list_return

                    else:
                        text = text.replace(text[index_begin:(index_fin+1)],
                                            client_params[index1][int(index2)])
                        while '\n' in text:
                            text = text.replace('\n', '')
                        list_return = [text]
                        return list_return

                elif doc_type == '2':
                    index1 = int(text[(index_begin+9):index_fin])
                    list_cell = []
                    for k in data[1:]:
                        if k is not None and k != '':
                            current_cell = k[index1].lower()
                            if current_cell[-1] != ';' and current_cell[-1] != '.':
                                current_cell = current_cell+';'
                            if current_cell[-1] == '.':
                                current_cell = current_cell[:-1]+';'
                            list_cell.append(current_cell)
                    list_cell = list(set(list_cell))
                    last_cell = list_cell.pop()
                    list_cell.append(last_cell[:-1]+'.')
                    list_return = [int(doc_type),
                                   list_cell]
                    return list_return

            list_return = [text]

            return list_return


        def doc_generator(sample: list, data: list, client_params: list) -> Document():
            doc = Document()
            section = doc.sections[0]

            if sample[1] is True:
                section.page_height = Cm(21)
                section.page_width = Cm(29.7)
                section.orientation = WD_ORIENT.LANDSCAPE
            else:
                section.page_height = Cm(29.7)
                section.page_width = Cm(21.0)
                section.orientation = WD_ORIENT.PORTRAIT

            section.left_margin = Mm(20)
            section.right_margin = Mm(20)
            section.top_margin = Mm(20)
            section.bottom_margin = Mm(20)
            section.header_distance = Mm(10)
            section.footer_distance = Mm(10)

            style = doc.styles['Normal']
            style.font.name = 'Times'
            style.font.size = Pt(14)
            style.paragraph_format.space_before = Mm(0)
            style.paragraph_format.space_after = Mm(0)
            style.paragraph_format.line_spacing = 1.0
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            style.paragraph_format.first_line_indent = Cm(1.25)

            if sample[2] == True:
                par = doc.add_paragraph('УТВЕРЖДАЮ')
                par_format = par.paragraph_format
                par_format.space_after = Mm(5)
                par_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                par = doc.add_paragraph(str(client_params[0][2])[0].title() +
                                        str(client_params[0][2])[1:] + ' ' +
                                        client_params[0][1])
                par_format = par.paragraph_format
                par_format.space_after = Mm(5)
                par_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                par_format.first_line_indent=Cm(0)
                if sample[1] == True:
                    par_format.left_indent=Cm(12.8)
                else:
                    par_format.left_indent=Cm(8.5)

                par = doc.add_paragraph('_______________ '+client_params[0][3])
                par_format = par.paragraph_format
                par_format.space_after = Mm(5)
                par_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                par_format.first_line_indent=Cm(0)
                par_format.left_indent=Cm(8.5)

                par = doc.add_paragraph('«_____»_______________20_____г.')
                par_format = par.paragraph_format
                par_format.space_after = Mm(5)
                par_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            if sample[3] != '':
                par = doc.add_paragraph(sample[3])
                par_format = par.paragraph_format
                par_format.space_after = Mm(5)
                par_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                par_format.first_line_indent=Cm(0)
                par_format.left_indent=Cm(8.5)

            insert_text = text_replacement(text=sample[4],
                                           data=data,
                                           client_params=client_params)
            par = doc.add_paragraph()
            par_format = par.paragraph_format
            par_format.space_after = Mm(5)
            par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par_format.first_line_indent = Cm(0)
            run = par.add_run(insert_text)
            run.font.bold = True

            if sample[0] == False:
                mark = 0
                for i in range(0, len(sample[5])):
                    for j in range(0, len(sample[5][i])):
                        if insert_text[0] == 1 and mark > 0:
                            mark -= 1
                            continue
                        insert_text = text_replacement(text=sample[5][i][j],
                                                       data=data,
                                                       client_params=client_params)
                        if j == 0 and len(sample[5])>1:
                            par = doc.add_paragraph(insert_text)
                            par_format = par.paragraph_format
                            par_format.space_before = Mm(5)
                            par_format.space_after = Mm(5)
                            par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            par_format.first_line_indent = Cm(0)
                        else:
                            if insert_text[0] == 1:
                                mark = len(insert_text[1])
                                for k in range(0, len(insert_text[1])):
                                    if insert_text[1][k] == 1:
                                        insert_sample = sample[5][i][j+k+1]
                                        if k == insert_text[2]:
                                            insert_sample = insert_sample[:-1]+'.'
                                        doc.add_paragraph(insert_sample)
                            elif insert_text[0] == 2:
                                for k in range(0, len(insert_text[1])):
                                    doc.add_paragraph(insert_text[1][k])
                            else:
                                doc.add_paragraph(insert_text[0])

            if sample[0] == True:
                sample_lenth = len(sample[5][0])

                table = doc.add_table(rows=1, cols=sample_lenth)
                table.style = 'Table Grid'
                style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                style.paragraph_format.first_line_indent = Cm(0)
                style.font.size = Pt(12)

                hdr_cells = table.rows[0].cells
                for j in range(0, sample_lenth):
                    hdr_cells[j].text = sample[5][0][j]

                for i in range(1, len(sample[5])):
                    row_cells = table.add_row().cells
                    for j in range(0, sample_lenth):
                        row_cells[j].text = sample[5][i][j]

            return doc

        data = load_data('users/table_old.json')
        data = add_data(data=data)

        client_params = load_data('users/client_params.json')

        dir_download = askdirectory()

        docs_list = load_data('templates/docs_list.json')

        for a in docs_list:
            if a[0] == 'data':
                sample = [True, True, True, '',
                          'Карта коррупционных рисков и мер по их '
                          'минимизации (устранению) <index_1_0_1>',
                          data]
            else:
                sample = load_data(f'templates/docs/{a[0]}')
            doc = doc_generator(sample=sample,
                                data=data,
                                client_params=client_params)
            doc.save(f'{dir_download}/{a[1]}.docx')

        root.destroy()

        return None


    def close_window():
        root.destroy()
        return None

    def text_0_0_func(event):
        client_params[0][0]=text_0_0.get('1.0', 'end')
        save_client_params()
        return None

    def text_0_1_func(event):
        client_params[0][1]=text_0_1.get('1.0', 'end')
        save_client_params()
        return None

    def text_0_2_func(event):
        client_params[0][2]=text_0_2.get('1.0', 'end')
        save_client_params()
        return None

    def text_0_3_func(event):
        client_params[0][3]=text_0_3.get('1.0', 'end')
        save_client_params()
        return None

    def text_1_0_func(event):
        client_params[1][0]=text_1_0.get('1.0', 'end')
        save_client_params()
        return None

    def text_1_1_func(event):
        client_params[1][1]=text_1_1.get('1.0', 'end')
        save_client_params()
        return None

    def text_1_2_func(event):
        client_params[1][2]=text_1_2.get('1.0', 'end')
        save_client_params()
        return None

    def win_acc():
        window_acc = Toplevel(root)
        window_acc.title('Изменения сохраняются автоматически')
        window_acc.geometry('+50+250')

        Label(window_acc, font=('Arial', 16),
              text='Функции комиссии по противодействию коррупции:').pack(**params2)

        check_acc_0 = IntVar(value=client_params[2][0])
        check_acc_1 = IntVar(value=client_params[2][1])
        check_acc_2 = IntVar(value=client_params[2][2])
        check_acc_3 = IntVar(value=client_params[2][3])
        check_acc_4 = IntVar(value=client_params[2][4])
        check_acc_5 = IntVar(value=client_params[2][5])

        def save_window_acc():
            client_params[2] = [
                check_acc_0.get(),
                check_acc_1.get(),
                check_acc_2.get(),
                check_acc_3.get(),
                check_acc_4.get(),
                check_acc_5.get()]
            save_client_params()

        Checkbutton(window_acc,
                    text='разработка и обсуждение антикоррупционной политики',
                    command=save_window_acc, font=('Arial', 12),
                    variable=check_acc_0).pack(**params2)
        Checkbutton(window_acc,
                    text='разработка и обсуждение кодекса этики '
                         'и служебного поведения работников',
                    command=save_window_acc, font=('Arial', 12),
                    variable=check_acc_1).pack(**params2)
        Checkbutton(window_acc,
                    text='разработка и обсуждение плана противодействия '
                         'коррупции и отчетов о его выполнении',
                    command=save_window_acc, font=('Arial', 12),
                    variable=check_acc_2).pack(**params2)
        Checkbutton(window_acc,
                    text='обсуждение вопросов антикоррупционной деятельности',
                    command=save_window_acc, font=('Arial', 12),
                    variable=check_acc_3).pack(**params2)
        Checkbutton(window_acc,
                    text='выработка рекомендаций о мерах по совершенствованию '
                         'антикоррупционной деятельности',
                    command=save_window_acc, font=('Arial', 12),
                    variable=check_acc_4).pack(**params2)
        Checkbutton(window_acc,
                    text='разработка и корректировка карты коррупционных '
                         'рисков и мер по их минимизации (устранению)',
                    command=save_window_acc, font=('Arial', 12),
                    variable=check_acc_5).pack(**params2)


    def win_cac():
        window_cac = Toplevel(root)
        window_cac.title('Изменения сохраняются автоматически')
        window_cac.geometry('+50+250')

        Label(window_cac, font=('Arial', 16),
        text='Функции комиссии по соблюдению требований к служебному (должностному) '
             'поведению и урегулированию конфликта интересов:').pack(**params2)

        check_cac_0 = IntVar(value=client_params[3][0])
        check_cac_1 = IntVar(value=client_params[3][1])
        check_cac_2 = IntVar(value=client_params[3][2])
        check_cac_3 = IntVar(value=client_params[3][3])

        def save_window_cac():
            client_params[3] = [
                check_cac_0.get(),
                check_cac_1.get(),
                check_cac_2.get(),
                check_cac_3.get()]
            save_client_params()

        Checkbutton(window_cac,
                    text='рассмотрение вопросов соблюдения требований к служебному '
                         '(должностному) поведению и урегулированию конфликта интересов',
                    command=save_window_cac, font=('Arial', 12),
                    variable=check_cac_0).pack(**params2)
        Checkbutton(window_cac,
                    text='выработка рекомендаций о мерах по предотвращению '
                         'и урегулированию конфликта интересов',
                    command=save_window_cac, font=('Arial', 12),
                    variable=check_cac_1).pack(**params2)
        Checkbutton(window_cac,
                    text='выработка рекомендаций о применении дисциплинарных '
                         'взысканий и об иных мерах реагирования на выявленные '
                         'коррупционные и иные правонарушения',
                    command=save_window_cac, font=('Arial', 12),
                    variable=check_cac_2).pack(**params2)
        Checkbutton(window_cac,
                    text='разъяснение положений кодекса этики '
                         'и служебного поведения работников',
                    command=save_window_cac, font=('Arial', 12),
                    variable=check_cac_3).pack(**params2)


    def win_act():
        window_act = Toplevel(root)
        window_act.title('Изменения сохраняются автоматически')
        window_act.geometry('+50+50')

        Label(window_act, font=('Arial', 16),
              text='Набор антикоррупционных инструментов:').pack(**params2)

        check_act_0 = IntVar(value=client_params[4][0])
        check_act_1 = IntVar(value=client_params[4][1])
        check_act_2 = IntVar(value=client_params[4][2])
        check_act_3 = IntVar(value=client_params[4][3])
        check_act_4 = IntVar(value=client_params[4][4])
        check_act_5 = IntVar(value=client_params[4][5])
        check_act_6 = IntVar(value=client_params[4][6])
        check_act_7 = IntVar(value=client_params[4][7])
        check_act_8 = IntVar(value=client_params[4][8])
        check_act_9 = IntVar(value=client_params[4][9])
        check_act_10 = IntVar(value=client_params[4][10])
        check_act_11 = IntVar(value=client_params[4][11])
        check_act_12 = IntVar(value=client_params[4][12])
        check_act_13 = IntVar(value=client_params[4][13])
        check_act_14 = IntVar(value=client_params[4][14])
        check_act_15 = IntVar(value=client_params[4][15])
        check_act_16 = IntVar(value=client_params[4][16])
        check_act_17 = IntVar(value=client_params[4][17])

        def save_window_act():
            client_params[4] = [
                check_act_0.get(),
                check_act_1.get(),
                check_act_2.get(),
                check_act_3.get(),
                check_act_4.get(),
                check_act_5.get(),
                check_act_6.get(),
                check_act_7.get(),
                check_act_8.get(),
                check_act_9.get(),
                check_act_10.get(),
                check_act_11.get(),
                check_act_12.get(),
                check_act_13.get(),
                check_act_14.get(),
                check_act_15.get(),
                check_act_16.get(),
                check_act_17.get()]
            save_client_params()

        Checkbutton(window_act,
                    text='определение должностного лица (структурного '
                         'подразделения) организации, ответственного за '
                         'профилактику коррупционных и иных правонарушений',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_0).pack(**params4)
        Checkbutton(window_act,
                    text='оценка коррупционных рисков и разработка '
                         'мер по их минимизации в организации',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_1).pack(**params4)
        Checkbutton(window_act,
                    text='разработка, принятие и внедрение стандартов и процедур, '
                         'направленных на обеспечение добросовестной работы организации',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_2).pack(**params4)
        Checkbutton(window_act,
                    text='разработка, принятие и внедрение кодекса '
                         'этики и служебного поведения работников',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_3).pack(**params4)
        Checkbutton(window_act,
                    text='разработка, принятие и внедрение порядка предотвращения '
                         'и урегулирования конфликта интересов работников',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_4).pack(**params4)
        Checkbutton(window_act,
                    text='недопущение составления неофициальной отчетности '
                         'и использования поддельных документов',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_5).pack(**params4)
        Checkbutton(window_act,
                    text='включение в договоры, связанные с хозяйственной деятельностью '
                         'организации, положений о соблюдении антикоррупционных '
                         'стандартов (антикоррупционной оговорки)',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_6).pack(**params4)
        Checkbutton(window_act,
                    text='введение в правила внутреннего трудового распорядка и '
                         'трудовые договоры работников антикоррупционных положений',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_7).pack(**params4)
        Checkbutton(window_act,
                    text='введение процедуры уведомления работником о случаях '
                         'склонения его к совершению коррупционных правонарушений',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_8).pack(**params4)
        Checkbutton(window_act,
                    text='введение процедуры уведомления работником о ставших ему '
                         'известными случаях совершения коррупционных правонарушений',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_9).pack(**params4)
        Checkbutton(window_act,
                    text='введение процедуры уведомления работником о возникшем '
                         'конфликте интересов либо о возможности его возникновения',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_10).pack(**params4)
        Checkbutton(window_act,
                    text='введение процедуры предотвращения и урегулирования '
                         'конфликта интересов работников',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_11).pack(**params4)
        Checkbutton(window_act,
                    text='введение процедур защиты работников, сообщивших '
                         'о коррупционных правонарушениях в деятельности организации',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_12).pack(**params4)
        Checkbutton(window_act,
                    text='ознакомление работников под роспись с локальными '
                         'нормативными актами, регламентирующими вопросы '
                         'предупреждения и противодействия коррупции в организации',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_13).pack(**params4)
        Checkbutton(window_act,
                    text='проведение обучающих мероприятий по вопросам '
                         'профилактики и противодействия коррупции',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_14).pack(**params4)
        Checkbutton(window_act,
                    text='организация информирования и консультирования работников '
                         'по вопросам применения (соблюдения) антикоррупционных '
                         'стандартов и процедур, исполнения обязанностей',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_15).pack(**params4)
        Checkbutton(window_act,
                    text='сотрудничество организации с контрольными '
                         '(надзорными) и правоохранительными органами',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_16).pack(**params4)
        Checkbutton(window_act,
                    text='подготовка и представление отчетов о проводимой работе '
                         'в сфере противодействия коррупции и достигнутых результатах',
                    command=save_window_act, font=('Arial', 12),
                    variable=check_act_17).pack(**params4)

    root.overrideredirect(True)
    root.state('zoomed')
    root.grid_columnconfigure(index=0, weight=1)
    root.grid_rowconfigure(index=0, weight=1)

    global client_params
    client_params = load_data('users/client_params.json')

    menu = Menu(root)
    root.config(menu=menu)
    menu.add_command(label='   = Выгрузить пакет документов =   ',
                     command=download_docs)
    menu.add_command(label='   = Выйти =   ',
                     command=close_window)

    params1 = {}
    params2 = {'anchor':'nw', 'padx':5, 'pady':10}
    params3 = {'anchor':'nw', 'padx':5, 'pady':0}
    params4 = {'anchor':'nw', 'padx':5, 'pady':5}

    Label(root, **params1, font=('Arial', 16),
          text='Заполните данные об организации').pack(**params2)

    Label(root, **params1, font=('Arial', 14),
          text='Название организации (в именительном падеже)').pack(**params2)
    text_0_0 = Text(root, width=135, height=2, font=('Arial', 14), wrap=WORD)
    text_0_0.pack(**params3)
    text_0_0.insert(index='1.0', chars=client_params[0][0])
    text_0_0.bind('<KeyRelease>', text_0_0_func)

    Label(root, **params1, font=('Arial', 14),
          text='Название организации (в родительном падеже)').pack(**params2)
    text_0_1 = Text(root, width=135, height=2, font=('Arial', 14), wrap=WORD)
    text_0_1.pack(**params3)
    text_0_1.insert(index='1.0', chars=client_params[0][1])
    text_0_1.bind('<KeyRelease>', text_0_1_func)

    Label(root, **params1, font=('Arial', 14),
          text='Должность руководителя организации '
               '(со строчной буквы)').pack(**params2)
    text_0_2 = Text(root, width=135, height=1, font=('Arial', 14), wrap=WORD)
    text_0_2.pack(**params3)
    text_0_2.insert(index='1.0', chars=client_params[0][2])
    text_0_2.bind('<KeyRelease>', text_0_2_func)

    Label(root, **params1, font=('Arial', 14),
          text='Инициалы и фамилия руководителя организации').pack(**params2)
    text_0_3 = Text(root, width=135, height=1, font=('Arial', 14), wrap=WORD)
    text_0_3.pack(**params3)
    text_0_3.insert(index='1.0', chars=client_params[0][3])
    text_0_3.bind('<KeyRelease>', text_0_3_func)


    Label(root, **params1, font=('Arial', 1)).pack(pady=5)
    Label(root, **params1, font=('Arial', 16),
          text='Заполните данные об ответственных '
               'за профилактику коррупции').pack(**params2)

    Label(root, **params1, font=('Arial', 14),
          text='Должность ответственного за профилактику коррупционных и '
               'иных правонарушений (со строчной буквы)').pack(**params2)
    text_1_0 = Text(root, width=135, height=1, font=('Arial', 14), wrap=WORD)
    text_1_0.pack(**params3)
    text_1_0.insert(index='1.0', chars=client_params[1][0])
    text_1_0.bind('<KeyRelease>', text_1_0_func)

    Label(root, **params1, font=('Arial', 14),
          text='Должность ответственного за соблюдение ограничений при заключении '
               'трудовых договоров (со строчной буквы)').pack(**params2)
    text_1_1 = Text(root, width=135, height=1, font=('Arial', 14), wrap=WORD)
    text_1_1.pack(**params3)
    text_1_1.insert(index='1.0', chars=client_params[1][1])
    text_1_1.bind('<KeyRelease>', text_1_1_func)

    Label(root, **params1, font=('Arial', 14),
          text='Должность ответственного за соблюдение ограничений при заключении '
               'гражданско-правовых договоров (со строчной буквы)').pack(**params2)
    text_1_2 = Text(root, width=135, height=1, font=('Arial', 14), wrap=WORD)
    text_1_2.pack(**params3)
    text_1_2.insert(index='1.0', chars=client_params[1][2])
    text_1_2.bind('<KeyRelease>', text_1_2_func)

    Button(text='Уточнить функции комиссии по противодействию коррупции',
           font=('Arial', 16), command=win_acc).pack(**params4)

    Button(text='Уточнить функции комиссии по соблюдению требований к служебному '
                '(должностному) поведению и урегулированию конфликта интересов',
           font=('Arial', 16), command=win_cac).pack(**params4)

    Button(text='Уточнить набор антикоррупционных инструментов',
           font=('Arial', 16), command=win_act).pack(**params4)

    root.mainloop()

    return None


def main() -> None:
    status, file_path_in = startpage()
    if status == 0:
        data = load_data(file_path_in)
        status = firstpage(data)
    if status == 0:
        secondpage()
    return None


if __name__ == '__main__':
    main()
